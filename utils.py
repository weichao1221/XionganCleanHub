import datetime
import json
import os
import re
import ssl
import stat
import subprocess
import sys
import urllib.request
from collections import defaultdict
from datetime import datetime

import appdirs
import requests
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from statics import StaticSource
from deepdiff import DeepDiff


class about_numbers():
    @staticmethod
    def number_to_chinese(num):

        chinese_number_map = {
            0: '零',
            1: '壹',
            2: '贰',
            3: '叁',
            4: '肆',
            5: '伍',
            6: '陆',
            7: '柒',
            8: '捌',
            9: '玖',
        }
        chinese_unit_map = {
            10: '拾',
            100: '佰',
            1000: '仟',
            10000: '万',
            100000000: '亿',
            1000000000000: '兆',
            10000000000000000: '京',
            100000000000000000000: '垓',
            1000000000000000000000000: '秭',
            10000000000000000000000000000: '穰',
            100000000000000000000000000000000: '沟',
        }

        def convert_chunk(chunk):
            result = ''
            length = len(chunk)
            for i, digit in enumerate(chunk):
                digit = int(digit)
                if digit != 0:
                    result += chinese_number_map[digit]
                    if i < length - 1:
                        unit = 10 ** (length - i - 1)
                        if unit in chinese_unit_map:
                            result += chinese_unit_map[unit]
                        elif unit % 1000 == 0:
                            result += '仟'
                        elif unit % 100 == 0:
                            result += '佰'
                        elif unit % 10 == 0:
                            result += '拾'
                else:
                    # 处理单独的零
                    if i < length - 1 and chunk[i + 1] != '0':
                        result += chinese_number_map[digit]

            return result

        def convert_integer_part(num_str):
            if num_str == '0':
                return chinese_number_map[0]

            length = len(num_str)

            # 处理整数部分
            result = ''
            if length <= 4:
                result = convert_chunk(num_str)
            elif length <= 8:
                result = convert_chunk(num_str[:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 12:
                result = convert_chunk(num_str[:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 16:
                result = convert_chunk(num_str[:length - 12]) + '兆' + convert_chunk(
                    num_str[length - 12:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 20:
                result = convert_chunk(num_str[:length - 16]) + '京' + convert_chunk(
                    num_str[length - 16:length - 12]) + '兆' + convert_chunk(
                    num_str[length - 12:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 24:
                result = convert_chunk(num_str[:length - 20]) + '垓' + convert_chunk(
                    num_str[length - 20:length - 16]) + '京' + convert_chunk(
                    num_str[length - 16:length - 12]) + '兆' + convert_chunk(
                    num_str[length - 12:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 28:
                result = convert_chunk(num_str[:length - 24]) + '秭' + convert_chunk(
                    num_str[length - 24:length - 20]) + '垓' + convert_chunk(
                    num_str[length - 20:length - 16]) + '京' + convert_chunk(
                    num_str[length - 16:length - 12]) + '兆' + convert_chunk(
                    num_str[length - 12:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 32:
                result = convert_chunk(num_str[:length - 28]) + '穰' + convert_chunk(
                    num_str[length - 28:length - 24]) + '秭' + convert_chunk(
                    num_str[length - 24:length - 20]) + '垓' + convert_chunk(
                    num_str[length - 20:length - 16]) + '京' + convert_chunk(
                    num_str[length - 16:length - 12]) + '兆' + convert_chunk(
                    num_str[length - 12:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            elif length <= 36:
                result = convert_chunk(num_str[:length - 32]) + '沟' + convert_chunk(
                    num_str[length - 32:length - 28]) + '穰' + convert_chunk(
                    num_str[length - 28:length - 24]) + '秭' + convert_chunk(
                    num_str[length - 24:length - 20]) + '垓' + convert_chunk(
                    num_str[length - 20:length - 16]) + '京' + convert_chunk(
                    num_str[length - 16:length - 12]) + '兆' + convert_chunk(
                    num_str[length - 12:length - 8]) + '亿' + convert_chunk(
                    num_str[length - 8:length - 4]) + '万' + convert_chunk(num_str[length - 4:])
            else:
                result = '超出范围'
            return result

        def convert_decimal_part(decimal_str):
            result = ''
            if not decimal_str or decimal_str == "0" or decimal_str == "00":
                result = "整"
            else:
                if int(decimal_part[0]) == 0:
                    result += "零"
                else:
                    result += chinese_number_map[int(decimal_part[0])] + "角"
                if len(decimal_part) > 1 and decimal_part[1] != "0":
                    result += chinese_number_map[int(decimal_part[1])] + "分"
                else:
                    result += "整"
            return result

        num_str = str(num)
        if '.' in num_str:
            integer_part, decimal_part = num_str.split('.')  # 分割整数和小数部分
            integer_chinese = convert_integer_part(integer_part)  # 转换整数部分
            decimal_chinese = convert_decimal_part(decimal_part)  # 转换小数部分

            if decimal_chinese:
                result = integer_chinese + '元' + decimal_chinese  # 拼接整数和小数部分
            else:
                result = integer_chinese + '元整'  # 拼接整数和小数部分
        else:
            result = convert_integer_part(num_str) + "元整"

        return result


class about_word():
    @staticmethod
    # 封面标题（大）项目名称
    def fengmian_doc1(doc, text_content: str, font_size: int):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        text = par.add_run(text_content)
        text.font.bold = True
        text.font.size = Pt(font_size)  # 三号字对应16磅
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")

    @staticmethod
    # 封面标题（大）项目名称
    def fengmian_doc_under_line(doc, text_content: str, font_size: int, underline: bool = True):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        text = par.add_run(f"{text_content}")
        text2 = par.add_run("工程")
        text.underline = underline
        text2.underline = False
        text.font.bold = True
        text2.font.bold = True
        text.font.size = Pt(font_size)  # 三号字对应16磅
        text2.font.size = Pt(font_size)
        text.font.name = "Times New Roman"
        text2.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        text2.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")

    @staticmethod
    def fengmian_doc3(doc, text_content: str):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        text = par.add_run(text_content)
        text.font.bold = True
        text.font.size = Pt(22)  # 三号字对应16磅
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")

    # 封面标题（小）编制单位、编制时间等
    @staticmethod
    def fengmian_doc2(doc, text_content: str):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        # par.paragraph_format.first_line_indent = par.style.font.size * 2
        text = par.add_run(text_content)
        text.font.bold = False
        text.font.size = Pt(16)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"方正小标宋简体")

    @staticmethod
    def fengmian_doc4(doc, text_content: str):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        # par.paragraph_format.first_line_indent = par.style.font.size * 2
        text = par.add_run(text_content)
        text.font.bold = False
        text.font.size = Pt(18)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    # 定义一级标题格式
    @staticmethod
    def Heading_1(doc, Heading_1: str):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        text = par.add_run(Heading_1)
        # text.font.bold = True
        par.paragraph_format.first_line_indent = par.style.font.size * 2
        text.font.size = Pt(16)  # 三号字对应16磅
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")

    # 定义二级标题格式
    @staticmethod
    def Heading_2(doc, Heading_2: str):
        Heading = doc.add_heading("", level=2)
        Heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        Heading.paragraph_format.line_spacing = 1
        Heading.paragraph_format.space_before = Pt(12)
        Heading.paragraph_format.space_after = Pt(12)
        # print(f"二级标题的首行缩进值{Heading.style.font.size, type(Heading.style.font.size)}")
        Heading.paragraph_format.first_line_indent = Heading.style.font.size * 2
        text_Heading = Heading.add_run(Heading_2)
        text_Heading.font.bold = True
        text_Heading.font.size = Pt(14)
        text_Heading.font.name = "Times New Roman"
        text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")
        text_Heading.font.color.rgb = RGBColor(0, 0, 0)

    # 定义三级标题格式
    @staticmethod
    def Heading_3(doc, Heading_3: str):
        Heading = doc.add_heading("", level=3)
        Heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        Heading.paragraph_format.line_spacing = 1
        Heading.paragraph_format.space_before = Pt(12)
        Heading.paragraph_format.space_after = Pt(12)
        Heading.style.font.size = Pt(14)
        # print(f"三级标题的首行缩进值{Heading.style.font.size, type(Heading.style.font.size)}")
        Heading.paragraph_format.first_line_indent = Heading.style.font.size * 2
        text_Heading = Heading.add_run(Heading_3)
        text_Heading.font.bold = False
        text_Heading.font.size = Pt(14)
        text_Heading.font.name = "Times New Roman"
        text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")
        text_Heading.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    # 自定义标题
    def Heading_union(doc, text_content: str, layout, font_name, font_size):
        Heading = doc.add_heading("", level=2)
        if layout == "左对齐":
            Heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif layout == "居中":
            Heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            Heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        Heading.paragraph_format.line_spacing = 1
        Heading.paragraph_format.space_before = Pt(12)
        Heading.paragraph_format.space_after = Pt(12)
        # Heading.paragraph_format.first_line_indent = Heading.style.font.size * 2
        text_Heading = Heading.add_run(text_content)
        text_Heading.font.bold = True
        text_Heading.font.size = Pt(font_size)
        text_Heading.font.name = "Times New Roman"
        # text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")
        text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), f'{font_name}')
        text_Heading.font.color.rgb = RGBColor(0, 0, 0)

    # 定义正文格式
    @staticmethod
    def Normal_doc(doc, text_content: str):
        # print(f"新增了一个正文{text_content}")
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.style.font.size = Pt(16)
        par.paragraph_format.first_line_indent = par.style.font.size * 2
        text = par.add_run(text_content)
        text.font.bold = False
        text.font.size = Pt(14)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    def Normal_doc_仿宋三号加粗(doc, text_content: str):
        # print(f"新增了一个正文{text_content}")
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.style.font.size = Pt(16)
        # par.paragraph_format.first_line_indent = par.style.font.size * 2
        text = par.add_run(text_content)
        text.font.bold = True
        text.font.size = Pt(16)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    def Normal_doc_red(doc, text_content: str):
        # print(f"新增了一个正文{text_content}")
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        text = par.add_run(text_content)
        text.font.bold = True
        # 添加字体颜色为红色、斜体
        text.font.color.rgb = RGBColor(255, 0, 0)
        text.font.italic = True  # 斜体
        text.font.size = Pt(16)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    # 不缩进正文
    def Normal_doc_not_suojin(doc, text_content: str):
        # print(f"新增了一个正文{text_content}")
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 做对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        text = par.add_run(text_content)
        text.font.bold = False  # 不加粗
        par.paragraph_format.first_line_indent = Pt(0)  # 首行不缩进
        # 添加字体颜色为红色、斜体
        # text.font.color.rgb = RGBColor(255, 0, 0)
        # text.font.italic = True # 斜体
        text.font.size = Pt(16)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    def add_header(doc, header_text):
        section = doc.sections[0]
        header = section.header

        # 添加段落（如果已有则使用第一个）
        par = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)

        # 添加文本
        run = par.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(12)  # 小四号（5号字体）
        run.font.name = "Times New Roman"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    def set_cell_format(cell, text_content: str):
        """设置单元格文字格式（仿宋_GB2312，12号，两端对齐）"""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format = cell.paragraphs[0].paragraph_format
        p_format.line_spacing = 1.5
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)

        run = cell.paragraphs[0].add_run(text_content)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    # 结尾落款
    @staticmethod
    def company_name(doc, company_name):
        # print(f"新增了一个公司名{company_name}")
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.first_line_indent = par.style.font.size * 2
        text = par.add_run(company_name)
        text.font.bold = False
        text.font.size = Pt(16)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    def created_time(doc):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 两端对齐
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.first_line_indent = par.style.font.size * 2
        text = par.add_run(f'{datetime.now().strftime("%Y年%m月%d日")}')
        text.font.bold = False
        text.font.size = Pt(14)
        text.font.name = "Times New Roman"
        text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    # 插入新的一页
    @staticmethod
    def insert_new_section(doc):
        new_section = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
        new_section.start_type = WD_SECTION_START.EVEN_PAGE

    @staticmethod
    def insert_img(doc, img_path, width):
        par = doc.add_paragraph("")
        par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = par.add_run()
        run.add_picture(img_path, width=Cm(width))

    @staticmethod
    def set_all_layout(doc):
        """
        设置整个文档的页边距（全文有效）
        上：2.54厘米 | 下：2.54厘米 | 左：3.17厘米 | 右：3.17厘米 | 装订线：0厘米
        """
        # 遍历所有节（Section）并设置页边距
        for section in doc.sections:
            # 设置页边距（单位：厘米）
            section.top_margin = Cm(2.54)  # 上边距 2.54cm
            section.bottom_margin = Cm(2.54)  # 下边距 2.54cm
            section.left_margin = Cm(3.17)  # 左边距 3.17cm
            section.right_margin = Cm(3.17)  # 右边距 3.17cm
            section.gutter = Cm(0)  # 装订线 0cm

    # 新增的测试性的内容

    # 进度款封面
    @staticmethod
    def add_cover_section_for_progress_payment(doc, project_name):
        time_now = datetime.now().strftime("%Y年%m月%d日")
        about_word.Normal_doc(doc, "")
        about_word.insert_img(doc, "./logo.jpg", 3.33)
        about_word.fengmian_doc1(doc, project_name, 22)
        about_word.Normal_doc(doc, "")
        about_word.Normal_doc(doc, "")
        about_word.fengmian_doc1(doc, f"控制价编制报告", 22)
        for i in range(8):
            about_word.Normal_doc(doc, "")
        about_word.fengmian_doc1(doc, "委托单位：中国雄安集团生态建设投资有限公司", 16)
        about_word.fengmian_doc1(doc, "编制单位：北京北咨工程咨询有限公司", 16)
        about_word.fengmian_doc1(doc, f"{time_now}", 16)

    @staticmethod
    def add_cover_section(doc, source_data, kind):
        time_now = datetime.now().strftime("%Y年%m月%d日")
        project_name = source_data.get("project_name", "测试项目名称")
        qishu = source_data.get("qishu", "测试")
        about_word.Normal_doc(doc, "")
        about_word.insert_img(doc, "./logo.jpg", 3.33)
        about_word.fengmian_doc1(doc, project_name, 22)
        about_word.Normal_doc(doc, "")
        about_word.Normal_doc(doc, "")
        about_word.fengmian_doc1(doc, f"{kind}审核报告", 22)
        about_word.fengmian_doc1(doc, f"{qishu}", 18)
        for i in range(8):
            about_word.Normal_doc(doc, "")
        about_word.fengmian_doc1(doc, "委托单位：中国雄安集团生态建设投资有限公司", 16)
        about_word.fengmian_doc1(doc, "编制单位：北京北咨工程咨询有限公司", 16)
        about_word.fengmian_doc1(doc, f"{time_now}", 16)

    @staticmethod
    def add_title_page(doc, header_text, logo_path):
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # 页眉段落（左对齐）
        header_par = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 插入 logo 图片
        run_img = header_par.add_run()
        run_img.add_picture(logo_path, width=Cm(1), height=Cm(1))

        # 插入页眉文字（与 logo 同行）
        run_text = header_par.add_run(f"    {header_text}")  # 加空格让文字稍微右移
        run_text.font.size = Pt(12)
        run_text.font.bold = True
        run_text.font.name = "Times New Roman"
        run_text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

        about_word.Normal_doc(doc, "")
        about_word.Normal_doc(doc, f"委托单位：")
        about_word.Normal_doc(doc, "报告编制单位：北京北咨工程咨询有限公司")
        about_word.Normal_doc(doc, "")

    @staticmethod
    def remove_first_page(doc):
        # 获取所有段落
        paragraphs = doc.paragraphs

        # 找到第一页的段落（通常是前几个段落）
        # 这里假设第一页段落数量不多，可以根据实际情况调整
        for i in range(min(5, len(paragraphs))):
            p = paragraphs[i]
            p.clear()  # 清空段落内容

    @staticmethod
    def add_toc_section(doc):
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.header.is_linked_to_previous = False

    @staticmethod
    def add_body_section(doc, header_text=None, footer_text=None, logo_path=None):
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # 页眉段落（左对齐）
        header_par = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        if logo_path:
            # 插入 logo 图片
            run_img = header_par.add_run()
            run_img.add_picture(logo_path, width=Cm(1), height=Cm(1))

        if header_text:
            # 插入页眉文字（与 logo 同行）
            run_text = header_par.add_run(f"    {header_text}")  # 加空格让文字稍微右移
            run_text.font.size = Pt(12)
            run_text.font.bold = True
            run_text.font.name = "Times New Roman"
            run_text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")
        if footer_text:
            # 页脚段落（居中）
            footer_par = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
            footer_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_footer = footer_par.add_run(footer_text)
            run_footer.font.size = Pt(12)
            run_footer.font.name = "Times New Roman"
            run_footer.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

    @staticmethod
    def add_appendix_section(doc, appendix_title="附录"):
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.header.is_linked_to_previous = False

    @staticmethod
    def reset_header_footer(section):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

    @staticmethod
    # 设置单元格垂直对齐方式的函数
    def set_cell_vertical_alignment(cell, align="top"):
        """
        设置单元格垂直对齐方式

        参数:
            cell: 单元格对象
            align: 对齐方式，可选值:
                - "top": 顶部对齐 (默认)
                - "center": 居中
                - "bottom": 底部对齐
        """
        # 获取单元格属性
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 创建垂直对齐元素
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), align)

        # 将垂直对齐设置添加到单元格属性中
        tcPr.append(v_align)

    @staticmethod
    # 写一个通用表格的方法
    def add_table(doc, headers, rows_data):
        cols = len(headers)
        rows = len(rows_data) + 1  # +1是表示标题行
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        # 第一列 1.25厘米，第二列3.25厘米，第三列1.75厘米，第四列3.25厘米，第五列2.5厘米，第六列3.25厘米
        # table.columns[0].width = Cm(1.25)
        # table.columns[1].width = Cm(3.25)
        # table.columns[2].width = Cm(1.75)
        # table.columns[3].width = Cm(3.25)
        # table.columns[4].width = Cm(2.5)
        # table.columns[5].width = Cm(3.25)

        # 设置每行高度为0.81厘米
        for row in table.rows:
            row.height = Cm(0.81)

        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        for row_num, row_data in enumerate(rows_data):
            for col_num, cell_data in enumerate(row_data):
                table.cell(row_num + 1, col_num).text = str(cell_data)

        # 设置表格主题部分样式 - 宋体、10号，居中, 垂直居中
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                cell.paragraphs[0].runs[0].element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].runs[0].font.bold = False
                about_word.set_cell_vertical_alignment(cell, "center")

        # 设置第一行（标题行）的格式 - 宋体、加粗，11号，居中
        for cell in table.rows[0].cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # cell.paragraphs[0].style = "Heading1"
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
            cell.paragraphs[0].runs[0].font.bold = True
            about_word.set_cell_vertical_alignment(cell, "center")


class functions:
    @staticmethod
    def get_platform():
        """
        获取当前操作系统平台
        :return:
        """
        if sys.platform.startswith('win'):
            return 'win'
        elif sys.platform.startswith('linux'):
            return 'linux'
        elif sys.platform.startswith('darwin'):
            return 'mac'
        else:
            return 'unknown'

    @staticmethod
    def get_base_info_path():
        # 应用名称和公司名称（可自定义）
        app_name = "Qingbiao"
        app_author = "Willchalighter"

        # 获取用户数据目录
        user_data_dir = appdirs.user_data_dir(app_name, app_author)
        os.makedirs(user_data_dir, exist_ok=True)  # 确保目录存在

        # base_info.json 文件路径
        base_info_path = os.path.join(user_data_dir, "base_info.json")
        return base_info_path

    @staticmethod
    def get_base_info_data():
        # 应用名称和公司名称（可自定义）
        app_name = "Qingbiao"
        app_author = "Willchalighter"

        # 获取用户数据目录
        user_data_dir = appdirs.user_data_dir(app_name, app_author)
        os.makedirs(user_data_dir, exist_ok=True)  # 确保目录存在

        # base_info.json 文件路径
        try:
            base_info_path = os.path.join(user_data_dir, "base_info.json")
            with open(base_info_path, "r", encoding="utf-8") as file:
                base_info = json.load(file)
            return base_info
        except FileNotFoundError as e:
            print(f"读取基本信息失败: {e}")
            return {}

    @staticmethod
    def get_new_version_download_url():
        platform_info = functions.get_platform()
        software_version = functions.get_latest_version()
        gitee_base_url = "https://gitee.com/weichao1221/xionganCleanHubDownload/releases/download"
        platform_name = ""
        if platform_info == "win":
            platform_name = f"xa_qingbiao_win.exe"
        elif platform_info == "mac":
            platform_name = f"xa_qingbiao_mac_apple_silicon.dmg"
        download_url = f"{gitee_base_url}/{software_version}/{platform_name}"
        return download_url

    @staticmethod
    def get_latest_version():
        tokens = StaticSource.get_gitee_token()
        url = "https://gitee.com/api/v5/repos/weichao1221/xionganCleanHubDownload/releases/latest"

        headers = {"Authorization": f"token {tokens}"}
        try:
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            data = response.json()
            return data.get("tag_name", "")
        except requests.RequestException:
            return ""

    @staticmethod
    def compare_version_numbers(current_version: str, latest_version: str) -> bool:
        """
        比较当前版本号与最新版本号，判断是否需要更新。
        版本号格式如: '1.0.0'
        如果 latest_version > current_version，则返回 True，否则返回 False。
        """

        def parse(v):
            return [int(x) for x in v.split('.')]

        cur = parse(current_version)
        latest = parse(latest_version)

        # 对齐版本长度，例如 1.0 vs 1.0.1
        max_len = max(len(cur), len(latest))
        cur += [0] * (max_len - len(cur))
        latest += [0] * (max_len - len(latest))

        return latest > cur

    @staticmethod
    def normalize_text(text):
        if not isinstance(text, str):
            return text

        text = re.sub(r'\s+', '', text.strip())  # 去所有空白

        trans_table = str.maketrans({
            '，': ',', '。': '.', '；': ';', '：': ':', '？': '?', '！': '!',
            '“': '"', '”': '"', '‘': "'", '’': "'", '（': '(', '）': ')', '【': '[', '】': ']'
        })
        text = text.translate(trans_table)

        # 清理连续相同标点
        text = re.sub(r'([，,.;:!?])\1+', r'\1', text)

        return text

    @staticmethod
    def clean_for_excel(text):
        """彻底清理文本，杜绝Excel报错"""
        if not isinstance(text, str):
            return ""

        # 1. 替换常见的Word/Excel实体
        text = text.replace('_x000D_', '\n')
        text = text.replace('_x005F_x000D_', '\n')

        # 2. 替换所有不可见控制字符（0x00-0x1F 除了 \t \n \r）
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

        # 3. 统一换行符（只保留 \n）
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # 4. 可选：去除首尾空白 + 限制长度（防止超格）
        text = text.strip()
        if len(text) > 32767:  # Excel 单元格最大长度
            text = text[:32767] + "...[截断]"

        return text

    @staticmethod
    def get_10_result(kzj_data, tb_datas, pianchalv, is_shebei: bool = False):
        _10_jieguo = {}

        for tb_data in tb_datas:
            if not tb_data['投标信息'].get("有效性", True):
                continue
            tb_name = tb_data['投标信息']['投标人']
            xiafulv = tb_data['投标信息'].get("下浮率", 0)
            _10_jieguo[tb_name] = {}

            for zb_dx, tb_dx in zip(kzj_data['单项工程'], tb_data['单项工程']):
                dx_name = zb_dx['名称']
                _10_jieguo[tb_name][dx_name] = {}

                for zb_dw, tb_dw in zip(zb_dx['单位工程'], tb_dx['单位工程']):
                    dw_name = zb_dw['名称']
                    _10_jieguo[tb_name][dx_name][dw_name] = {}

                    # 分部清单
                    for zb_qd, tb_qd in zip(zb_dw['分部清单'], tb_dw['分部清单']):
                        if zb_qd['编码'] != tb_qd['编码']:
                            continue

                        zb_dj = float(zb_qd.get('综合单价_含设备' if is_shebei else '综合单价'))
                        tb_dj = float(tb_qd.get("综合单价_含设备" if is_shebei else "综合单价"))
                        if zb_dj == 0 or tb_dj == 0:
                            continue

                        base_price = zb_dj * (xiafulv / 100)
                        if base_price == 0:
                            continue

                        pianchalv_calc = (tb_dj - base_price) / base_price * 100
                        bucket = (f'大于{pianchalv}%' if pianchalv_calc > pianchalv else
                                  f'小于-{pianchalv}%' if pianchalv_calc < -pianchalv else None)
                        if not bucket:
                            continue

                        item = {
                            "编码": zb_qd['编码'],
                            "名称": tb_qd.get('名称', ''),
                            "项目特征": tb_qd.get('项目特征', ''),
                            "单位": tb_qd.get('单位', ''),
                            "工程量": tb_qd.get('数量', ''),
                            "控制价单价": zb_dj,
                            "下浮率": f"{xiafulv:.2f}%",
                            "基准价": round(base_price, 4),
                            "投标单价": tb_dj,
                            "偏差率": round(pianchalv_calc, 2),
                        }

                        _10_jieguo[tb_name][dx_name][dw_name].setdefault(bucket, []).append(item)

                    # 措施清单
                    for zb_cs, tb_cs in zip(zb_dw['措施清单'], tb_dw['措施清单']):
                        if zb_cs['编码'] != tb_cs['编码']:
                            continue

                        zb_dj = float(zb_cs.get('综合单价'))
                        tb_dj = float(tb_cs.get("综合单价"))
                        if zb_dj == 0 or tb_dj == 0:
                            continue

                        base_price = zb_dj * (xiafulv / 100)
                        if base_price == 0:
                            continue

                        pianchalv_calc = (tb_dj - base_price) / base_price * 100
                        bucket = (f'大于{pianchalv}%' if pianchalv_calc > pianchalv else
                                  f'小于-{pianchalv}%' if pianchalv_calc < -pianchalv else None)
                        if not bucket:
                            continue

                        item = {
                            "编码": zb_cs['编码'],
                            "名称": tb_cs.get('名称', ''),
                            "项目特征": tb_cs.get('项目特征', ''),
                            "单位": tb_cs.get('单位', ''),
                            "工程量": tb_cs.get('数量', ''),
                            "控制价单价": zb_dj,
                            "下浮率": f"{xiafulv:.2f}%",
                            "基准价": round(base_price, 4),
                            "投标单价": tb_dj,
                            "偏差率": round(pianchalv_calc, 2),
                        }

                        _10_jieguo[tb_name][dx_name][dw_name].setdefault(bucket, []).append(item)

        return _10_jieguo

    @staticmethod
    def get_10_result_and_qingdan_result(kzj_data, tb_datas, pianchalv, is_shebei: bool = False):
        qingdan_jieguo = {}
        _10_jieguo = {}
        for tb_data in tb_datas:
            if not tb_data['投标信息'].get("有效性", True):
                continue
            tb_name = tb_data['投标信息']['投标人']
            xiafulv = tb_data['投标信息'].get("下浮率", 0)  # 修正键名
            qingdan_jieguo[tb_name] = {}
            _10_jieguo[tb_name] = {}

            for zb_dx, tb_dx in zip(kzj_data['单项工程'], tb_data['单项工程']):
                dx_name = zb_dx['名称']
                qingdan_jieguo[tb_name][dx_name] = {}
                _10_jieguo[tb_name][dx_name] = {}

                for zb_dw, tb_dw in zip(zb_dx['单位工程'], tb_dx['单位工程']):
                    dw_name = zb_dw['名称']
                    qingdan_jieguo[tb_name][dx_name][dw_name] = {}
                    _10_jieguo[tb_name][dx_name][dw_name] = {}

                    for zb_qd, tb_qd in zip(zb_dw['分部清单'], tb_dw['分部清单']):
                        if zb_qd['编码'] != tb_qd['编码']:
                            continue
                        qingdan_jieguo[tb_name][dx_name][dw_name]['分部清单'] = {}
                        zb_name = functions.normalize_text(zb_qd.get('名称', ''))
                        tb_name_norm = functions.normalize_text(tb_qd.get('名称', ''))
                        zb_feat = functions.normalize_text(zb_qd.get('项目特征', ''))
                        tb_feat = functions.normalize_text(tb_qd.get('项目特征', ''))
                        zb_unit = functions.normalize_text(zb_qd.get('单位', ''))
                        tb_unit = functions.normalize_text(tb_qd.get('单位', ''))
                        zb_qty = functions.normalize_text(zb_qd.get('数量', ''))
                        tb_qty = functions.normalize_text(tb_qd.get('数量', ''))

                        if (zb_name != tb_name_norm or
                                zb_feat != tb_feat or
                                zb_unit != tb_unit or
                                float(zb_qty) != float(tb_qty)):
                            qingdan_jieguo[tb_name][dx_name][dw_name]['分部清单'] = {
                                "招标清单": zb_qd,
                                "投标清单": tb_qd
                            }
                        # 计算下浮后的单价 = 控制价单价 * 下浮率
                        if is_shebei:
                            zb_dj = float(zb_qd.get('综合单价_含设备'))
                            tb_dj = float(tb_qd.get("综合单价_含设备"))
                        else:
                            zb_dj = float(zb_qd.get('综合单价'))
                            tb_dj = float(tb_qd.get("综合单价"))
                        if zb_dj == 0 or tb_dj == 0:
                            continue  # 跳过控制价为0
                        base_price = zb_dj * (xiafulv / 100)  # 控制价下浮后基准价
                        if base_price == 0:
                            continue  # 避免除零

                        pianchalv_calc = (tb_dj - base_price) / base_price * 100
                        item = {
                            "编码": zb_qd['编码'],
                            "名称": tb_qd.get('名称', ''),
                            "项目特征": tb_qd.get('项目特征', ''),
                            "单位": tb_qd.get('单位', ''),
                            "工程量": tb_qd.get('数量', ''),
                            "控制价单价": zb_dj,
                            "下浮率": f"{xiafulv:.2f}%",
                            "基准价": round(base_price, 4),
                            "投标单价": tb_dj,
                            "偏差率": round(pianchalv_calc, 2),
                        }

                        # 按偏差率分桶
                        if pianchalv_calc > pianchalv:
                            bucket = f'大于{pianchalv}%'
                        elif pianchalv_calc < -pianchalv:
                            bucket = f'小于-{pianchalv}%'
                        else:
                            bucket = None

                        if bucket:
                            if bucket not in _10_jieguo[tb_name][dx_name][dw_name]:
                                _10_jieguo[tb_name][dx_name][dw_name][bucket] = []
                            _10_jieguo[tb_name][dx_name][dw_name][bucket].append(item)

                    for zb_cs, tb_cs in zip(zb_dw['措施清单'], tb_dw['措施清单']):
                        if zb_cs['编码'] != tb_cs['编码']:
                            continue
                        qingdan_jieguo[tb_name][dx_name][dw_name]['措施清单'] = {}

                        zb_name = functions.normalize_text(zb_cs.get('名称', ''))
                        tb_name_norm = functions.normalize_text(tb_cs.get('名称', ''))
                        zb_feat = functions.normalize_text(zb_cs.get('项目特征', ''))
                        tb_feat = functions.normalize_text(tb_cs.get('项目特征', ''))
                        zb_unit = functions.normalize_text(zb_cs.get('单位', ''))
                        tb_unit = functions.normalize_text(tb_cs.get('单位', ''))
                        zb_qty = functions.normalize_text(zb_cs.get('数量', ''))
                        tb_qty = functions.normalize_text(tb_cs.get('数量', ''))

                        if (zb_name != tb_name_norm or
                                zb_feat != tb_feat or
                                zb_unit != tb_unit or
                                float(zb_qty) != float(tb_qty)):
                            qingdan_jieguo[tb_name][dx_name][dw_name]['措施清单'] = {
                                "招标清单": zb_cs,
                                "投标清单": tb_cs
                            }

                        # 计算下浮后的单价 = 控制价单价 * 下浮率
                        zb_dj = float(zb_cs.get('综合单价'))
                        tb_dj = float(tb_cs.get("综合单价"))

                        if zb_dj == 0 or tb_dj == 0:
                            continue
                        base_price = zb_dj * (xiafulv / 100)  # 控制价下浮后基准价
                        if base_price == 0:
                            continue

                        pianchalv_calc = (tb_dj - base_price) / base_price * 100
                        item = {
                            "编码": zb_cs['编码'],
                            "名称": tb_cs.get('名称', ''),
                            "项目特征": tb_cs.get('项目特征', ''),
                            "单位": tb_cs.get('单位', ''),
                            "工程量": tb_cs.get('数量', ''),
                            "控制价单价": zb_dj,
                            "下浮率": f"{xiafulv:.2f}%",
                            "基准价": round(base_price, 4),
                            "投标单价": tb_dj,
                            "偏差率": round(pianchalv_calc, 2),
                        }
                        if pianchalv_calc > pianchalv:
                            bucket = f'大于{pianchalv}%'
                        elif pianchalv_calc < -pianchalv:
                            bucket = f'小于-{pianchalv}%'
                        else:
                            bucket = None

                        if bucket:
                            if bucket not in _10_jieguo[tb_name][dx_name][dw_name]:
                                _10_jieguo[tb_name][dx_name][dw_name][bucket] = []
                            _10_jieguo[tb_name][dx_name][dw_name][bucket].append(item)

                    for (key1, value1), (key2, value2) in zip(zb_dw['其他项目'].items(), tb_dw['其他项目'].items()):
                        qingdan_jieguo[tb_name][dx_name][dw_name]['其他项目'] = {}
                        diff = DeepDiff(value1, value2)
                        if diff:
                            qingdan_jieguo[tb_name][dx_name][dw_name]['其他项目'][key1] = {
                                "招标清单": value1,
                                "投标清单": value2
                            }
        return _10_jieguo, qingdan_jieguo

    @staticmethod
    def get_result(kzj_data, tb_data):
        result = []
        # 开始尝试对比控制价单项工程和投标单位2的单项工程中的内容
        for kzj_dx, tb_dx in zip(kzj_data['单项工程'], tb_data['单项工程']):
            for kzj_dw, tb_dw in zip(kzj_dx['单位工程'], tb_dx['单位工程']):
                for kzj_qd, tb_qd in zip(kzj_dw['分部清单'], tb_dw['分部清单']):
                    if kzj_qd['编码'] != tb_qd['编码']:
                        continue

                    kzj_duibi = {
                        '编码': kzj_qd['编码'],
                        '名称': functions.normalize_text(kzj_qd['名称']),
                        '项目特征': functions.normalize_text(kzj_qd['项目特征']),
                        '单位': functions.normalize_text(kzj_qd['单位']),
                        '数量': float(kzj_qd['数量']),
                    }
                    tb_duibi = {
                        '编码': tb_qd['编码'],
                        '名称': functions.normalize_text(tb_qd['名称']),
                        '项目特征': functions.normalize_text(tb_qd['项目特征']),
                        '单位': functions.normalize_text(tb_qd['单位']),
                        '数量': float(tb_qd['数量']),
                    }
                    df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                    if df:
                        data_result = {
                            # "公司名称": tb_data['投标信息']['投标人'],
                            "单项工程": kzj_dx['名称'],
                            "单位工程": kzj_dw['名称'],
                            "清单类别": "分部清单",
                            "招标清单": kzj_duibi,
                            "投标清单": tb_duibi,
                            # "差异结果": df,
                        }
                        result.append(data_result)

                # 措施项目
                for kzj_cs, tb_cs in zip(kzj_dw['措施清单'], tb_dw['措施清单']):
                    if kzj_cs['编码'] != tb_cs['编码']:
                        continue

                    kzj_duibi = {
                        '编码': kzj_cs['编码'],
                        '名称': functions.normalize_text(kzj_cs['名称']),
                        '项目特征': functions.normalize_text(kzj_cs['项目特征']),
                        '单位': functions.normalize_text(kzj_cs['单位']),
                        '数量': float(kzj_cs['数量']),
                    }
                    tb_duibi = {
                        '编码': tb_cs['编码'],
                        '名称': functions.normalize_text(tb_cs['名称']),
                        '项目特征': functions.normalize_text(tb_cs['项目特征']),
                        '单位': functions.normalize_text(tb_cs['单位']),
                        '数量': float(tb_cs['数量']),
                    }
                    df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                    if df:
                        data_result = {
                            "单项工程": kzj_dx['名称'],
                            "单位工程": kzj_dw['名称'],
                            "清单类别": "措施清单",
                            "招标清单": kzj_duibi,
                            "投标清单": tb_duibi,
                            # "差异结果": df,
                        }
                        result.append(data_result)
                for (key1, value1), (key2, value2) in zip(kzj_dw["其他项目"].items(), tb_dw['其他项目'].items()):
                    if key1 == "暂列金额明细":
                        for item1, item2 in zip(value1, value2):
                            kzj_duibi = {
                                "项目名称": item1['项目名称'],
                                "计量单位": item1['计量单位'],
                                "暂定金额": item1['暂定金额'],
                            }
                            tb_duibi = {
                                "项目名称": item2['项目名称'],
                                "计量单位": item2['计量单位'],
                                "暂定金额": item2['暂定金额'],
                            }
                            df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                            if df:
                                data_result = {
                                    "单项工程": kzj_dx['名称'],
                                    "单位工程": kzj_dw['名称'],
                                    "清单类别": "暂列金额明细",
                                    "招标清单": kzj_duibi,
                                    "投标清单": tb_duibi,
                                }
                                result.append(data_result)
                    elif key1 == "暂估价材料明细":
                        for item1, item2 in zip(value1, value2):
                            kzj_duibi = {
                                "招标材料号": item1['招标材料号'],
                                "材料名称": item1['材料名称'],
                                "规格型号": item1['规格型号'],
                                "计量单位": item1['计量单位'],
                                "数量": item1['数量'],
                                "暂定价": item1['暂定价'],
                            }
                            tb_duibi = {
                                "招标材料号": item2['招标材料号'],
                                "材料名称": item2['材料名称'],
                                "规格型号": item2['规格型号'],
                                "计量单位": item2['计量单位'],
                                "数量": item2['数量'],
                                "暂定价": item2['暂定价'],
                            }
                            df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                            if df:
                                data_result = {
                                    "单项工程": kzj_dx['名称'],
                                    "单位工程": kzj_dw['名称'],
                                    "清单类别": "暂估价材料明细",
                                    "招标清单": kzj_duibi,
                                    "投标清单": tb_duibi,
                                }
                                result.append(data_result)
                    elif key1 == "暂估价设备明细":
                        for item1, item2 in zip(value1, value2):
                            kzj_duibi = {
                                "招标材料号": item1['招标材料号'],
                                "材料名称": item1['材料名称'],
                                "规格型号": item1['规格型号'],
                                "计量单位": item1['计量单位'],
                                "数量": item1['数量'],
                                "暂定价": item1['暂定价'],
                            }
                            tb_duibi = {
                                "招标材料号": item2['招标材料号'],
                                "材料名称": item2['材料名称'],
                                "规格型号": item2['规格型号'],
                                "计量单位": item2['计量单位'],
                                "数量": item2['数量'],
                                "暂定价": item2['暂定价'],
                            }
                            df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                            if df:
                                data_result = {
                                    "单项工程": kzj_dx['名称'],
                                    "单位工程": kzj_dw['名称'],
                                    "清单类别": "暂估价设备明细",
                                    "招标清单": kzj_duibi,
                                    "投标清单": tb_duibi,
                                }
                                result.append(data_result)

                    elif key1 == "专业工程暂估明细":
                        for item1, item2 in zip(value1, value2):
                            print(item1, item2)
                            kzj_duibi = {
                                "工程名称": item1['工程名称'],
                                "工程内容": item1['工程内容'],
                                "金额": item1.get("金额", item1.get("含税金额", "获取数据出错")),
                            }
                            tb_duibi = {
                                "工程名称": item2['工程名称'],
                                "工程内容": item2['工程内容'],
                                "金额": item2.get("金额", item2.get("含税金额", "获取数据出错")),
                            }
                            df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                            if df:
                                data_result = {
                                    "单项工程": kzj_dx['名称'],
                                    "单位工程": kzj_dw['名称'],
                                    "清单类别": "专业工程暂估明细",
                                    "招标清单": kzj_duibi,
                                    "投标清单": tb_duibi,
                                }
                                result.append(data_result)
                    elif key1 == "计日工项":
                        for item1, item2 in zip(value1, value2):
                            kzj_duibi = {
                                "名称": item1['名称'],
                                "单位": item1['单位'],
                                "型号规格": item1['型号规格'],
                                "暂定数量": item1['暂定数量'],
                                "综合单价": item1['综合单价'],
                                "综合合价": item1['综合合价'],
                            }
                            tb_duibi = {
                                "名称": item2['名称'],
                                "单位": item2['单位'],
                                "型号规格": item2['型号规格'],
                                "暂定数量": item2['暂定数量'],
                                "综合单价": item2['综合单价'],
                                "综合合价": item2['综合合价'],
                            }
                            df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                            if df:
                                data_result = {
                                    "单项工程": kzj_dx['名称'],
                                    "单位工程": kzj_dw['名称'],
                                    "清单类别": "计日工项",
                                    "招标清单": kzj_duibi,
                                    "投标清单": tb_duibi,
                                }
                                result.append(data_result)
                    elif key1 == "总承包服务费项":
                        for item1, item2 in zip(value1, value2):
                            print(item1, item2)
                            kzj_duibi = {
                                "项目名称": item1['项目名称'],
                                "服务内容": item1['服务内容'],
                                "项目价值": item1['项目价值'],
                                "金额": item1['金额']
                            }
                            tb_duibi = {
                                "项目名称": item2['项目名称'],
                                "服务内容": item2['服务内容'],
                                "项目价值": item2['项目价值'],
                                "金额": item2['金额']
                            }
                            df = DeepDiff(kzj_duibi, tb_duibi, ignore_order=True)
                            if df:
                                data_result = {
                                    "单项工程": kzj_dx['名称'],
                                    "单位工程": kzj_dw['名称'],
                                    "清单类别": "总承包服务费项",
                                    "招标清单": kzj_duibi,
                                    "投标清单": tb_duibi,
                                }
                                result.append(data_result)
        return result

    @staticmethod
    def get_qingdan_result(kzj_data, tb_datas):
        qingdan_result = {}
        for tb_data in tb_datas:
            name = tb_data['投标信息']['投标人']
            result = functions.get_result(kzj_data, tb_data)
            print(name, result)
            if name not in qingdan_result:
                if len(result) == 0:    # 如果没有，则跳过
                    continue
                qingdan_result[name] = result
        return qingdan_result


    @staticmethod
    def get_jiamisuo_jieguo(tb_data):
        lock_to_companies = defaultdict(list)
        mac_to_companies = defaultdict(list)
        jiamisuo_jieguio = []

        # 第一步：收集所有锁和MAC
        for tb in tb_data:
            if not tb['投标信息'].get("有效性", True):
                continue
            company = tb['投标信息']['投标人']
            lock_str = tb['投标信息'].get('加密锁号_解密', '')
            mac_str = tb['投标信息'].get('MAC地址_解密', '')

            locks = [l.strip() for l in lock_str.replace('_x000D_', ',').split(',') if l.strip()]
            macs = [m.strip() for m in mac_str.replace('_x000D_', ',').split(',') if m.strip()]

            for l in locks:
                lock_to_companies[l].append(company)
            for m in macs:
                mac_to_companies[m].append(company)

        # 第二步：生成表格行
        for idx, tb in enumerate(tb_data, 1):
            if not tb['投标信息'].get("有效性", True):
                continue
            company = tb['投标信息']['投标人']
            total_price = tb['投标信息'].get('投标总价', '')
            lock_str = tb['投标信息'].get('加密锁号_解密', '')
            mac_str = tb['投标信息'].get('MAC地址_解密', '')

            locks = [l.strip() for l in lock_str.replace('_x000D_', ',').split(',') if l.strip()]
            macs = [m.strip() for m in mac_str.replace('_x000D_', ',').split(',') if m.strip()]

            lock_display = ', '.join(locks) if locks else '无'
            mac_display = ', '.join(macs) if macs else '无'

            # 重复锁
            dup_locks = [l for l in set(locks) if len(lock_to_companies[l]) > 1]
            dup_lock_companies = set()
            for l in dup_locks:
                others = [c for c in lock_to_companies[l] if c != company]
                dup_lock_companies.update(others)
            dup_lock_str = ', '.join(dup_lock_companies) if dup_lock_companies else '无'

            # 重复MAC
            dup_macs = [m for m in set(macs) if len(mac_to_companies[m]) > 1]
            dup_mac_companies = set()
            for m in dup_macs:
                others = [c for c in mac_to_companies[m] if c != company]
                dup_mac_companies.update(others)
            dup_mac_str = ', '.join(dup_mac_companies) if dup_mac_companies else '无'

            row = {
                "序号": str(idx),
                "投标人": company,
                "总价": str(total_price),
                "加密锁数量": str(len(locks)),
                "锁号": lock_display,
                "重复单位(锁)": dup_lock_str,
                "MAC数量": str(len(macs)),
                "物理地址": mac_display,
                "重复单位(MAC)": dup_mac_str,
            }
            jiamisuo_jieguio.append(row)

        return jiamisuo_jieguio

    @staticmethod
    def get_zongjia_jieguo(kzj_data, tb_datas):
        zongjia_jieguo = []
        kzj_jine = float(kzj_data['招标控制价信息']['控制价总价'])
        for tb_data in tb_datas:
            tb_name = tb_data['投标信息']['投标人']
            tb_jine = float(tb_data['投标信息']['投标总价'])
            rate = tb_data['投标信息']['下浮率']
            if rate > 100:
                xianjia = "超出"
            else:
                xianjia = "未超出"
            a_list = [tb_name, str(tb_jine), str(kzj_jine), f"{rate:.3f}", xianjia]
            zongjia_jieguo.append(a_list)
        return zongjia_jieguo

    @staticmethod
    def get_fuzhi_and_zero_jieguo(tb_datas):
        fuzhi_jieguo = {}
        zero_jieguo = {}
        for tb_data in tb_datas:
            if not tb_data['投标信息'].get("有效性", True):
                continue
            tb_name = tb_data['投标信息']['投标人']
            fuzhi_jieguo[tb_name] = {}
            zero_jieguo[tb_name] = {}
            for dx in tb_data['单项工程']:
                fuzhi_jieguo[tb_name][dx['名称']] = {}
                zero_jieguo[tb_name][dx['名称']] = {}
                for dw in dx['单位工程']:
                    fuzhi_jieguo[tb_name][dx['名称']][dw['名称']] = {}
                    zero_jieguo[tb_name][dx['名称']][dw['名称']] = {}
                    for qd in dw['分部清单']:
                        fuzhi_jieguo[tb_name][dx['名称']][dw['名称']]['分部清单'] = []
                        zero_jieguo[tb_name][dx['名称']][dw['名称']]['分部清单'] = []
                        if float(qd['综合单价']) < 0:
                            fuzhi_jieguo[tb_name][dx['名称']][dw['名称']]['分部清单'].append(qd)
                        if float(qd['综合单价']) == 0:
                            zero_jieguo[tb_name][dx['名称']][dw['名称']]['分部清单'].append(qd)
                    for cs in dw['措施清单']:
                        fuzhi_jieguo[tb_name][dx['名称']][dw['名称']]['措施清单'] = []
                        zero_jieguo[tb_name][dx['名称']][dw['名称']]['措施清单'] = []
                        if float(cs['综合单价']) < 0:
                            fuzhi_jieguo[tb_name][dx['名称']][dw['名称']]['措施清单'].append(cs)
                        if float(cs['综合单价']) == 0:
                            zero_jieguo[tb_name][dx['名称']][dw['名称']]['措施清单'].append(cs)

        return fuzhi_jieguo, zero_jieguo

    @staticmethod
    def get_update_temp_dir():
        app_name = "Qingbiao"
        app_author = "Willchalighter"
        user_data_dir = appdirs.user_data_dir(app_name, app_author)
        temp_dir = os.path.join(user_data_dir, "update_temp")
        os.makedirs(temp_dir, exist_ok=True)
        return temp_dir

    @staticmethod
    def auto_update_mac_app(dmg_url):
        tokens = "3d500aabc2e64f9f80df3eb9fe916712"
        download_dir = r"/Users/chao/Downloads"
        headers = {"Authorization": f"token {tokens}"}

        try:
            # 创建不验证 SSL 的上下文
            ssl_context = ssl.create_default_context()
            ssl_context.check_hostname = False
            ssl_context.verify_mode = ssl.CERT_NONE

            # 创建请求对象
            req = urllib.request.Request(dmg_url, headers=headers)

            # 使用 SSL 上下文下载
            response = urllib.request.urlopen(req, context=ssl_context)

            # 获取文件总大小
            total_size = int(response.headers.get('content-length', 0))
            filename = os.path.join(download_dir, os.path.basename(dmg_url))

            downloaded = 0
            block_size = 8192  # 8KB

            with open(filename, 'wb') as f:
                while True:
                    chunk = response.read(block_size)
                    if not chunk:
                        break
                    f.write(chunk)
                    downloaded += len(chunk)

                    # 显示进度条
                    if total_size > 0:
                        percent = downloaded / total_size * 100
                        bar_length = 50
                        filled_length = int(bar_length * downloaded // total_size)
                        bar = '█' * filled_length + '-' * (bar_length - filled_length)
                        sys.stdout.write(f'\r下载进度: |{bar}| {percent:.1f}% ({downloaded}/{total_size})')
                        sys.stdout.flush()

            print("\n下载完成！")
            print(f"文件已下载到: {filename}")

            # 写一个脚本，保存至目标文件夹
            sh_text = """#!/bin/bash

    # 简化版安装脚本
    APP_NAME="雄安清标.app"
    DMG_FILE="/Users/chao/Downloads/xa_qingbiao_mac_apple_silicon.dmg"
    MOUNT_POINT="/Volumes/雄安清标"

    echo "开始安装应用程序..."

    # 挂载 DMG
    hdiutil attach "$DMG_FILE" -nobrowse -mountpoint "$MOUNT_POINT"

    # 停止正在运行的应用
    echo "停止正在运行的应用程序..."
    pkill -f "雄安清标" || true
    sleep 2

    # 确保应用程序完全停止
    pkill -9 -f "雄安清标" || true
    sleep 1

    # 删除旧版本
    echo "删除旧版本应用程序..."
    rm -rf "/Applications/$APP_NAME"

    # 复制新版本
    echo "安装新版本应用程序..."
    cp -R "$MOUNT_POINT/$APP_NAME" "/Applications/"

    # 卸载 DMG
    echo "清理临时文件..."
    hdiutil detach "$MOUNT_POINT" -force

    # 修复应用程序权限
    echo "修复应用程序权限..."
    chmod -R 755 "/Applications/$APP_NAME"

    echo "安装完成！"

    # 可选：启动应用程序
    read -p "是否立即启动应用程序? (Y/n): " -n 1 -r
    echo
    if [[ $REPLY =~ ^[Yy]$ ]] || [[ -z "$REPLY" ]]; then
        echo "启动应用程序..."
        open "/Applications/$APP_NAME"
    fi
    """
            script_path = os.path.join(download_dir, "install.sh")
            with open(script_path, "w", encoding="utf-8") as file:
                file.write(sh_text)
            print(f"已生成安装脚本: {script_path}")

            # 在执行脚本前，确保已有权限
            print("设置脚本执行权限...")
            os.chmod(script_path, stat.S_IRWXU | stat.S_IRGRP | stat.S_IXGRP | stat.S_IROTH | stat.S_IXOTH)

            # 如果程序没退出，则等待程序退出后开始执行脚本
            print("等待当前程序退出...")

            # 检查当前是否有同名进程在运行
            current_pid = os.getpid()
            print(f"当前进程ID: {current_pid}")


            try:
                # 执行安装脚本
                result = subprocess.run(['/bin/bash', script_path],
                                        capture_output=False,
                                        text=True,
                                        check=True)

                print("-" * 50)
                print("安装脚本执行完成！")

                os.remove(script_path)
                print("安装脚本已删除")

            except subprocess.CalledProcessError as e:
                print(f"安装脚本执行失败: {e}")
                print("请手动执行安装脚本完成安装")
            except Exception as e:
                print(f"执行安装脚本时发生错误: {e}")
                print(f"请手动运行: bash {script_path}")

            print("退出程序...")
            sys.exit(0)


        except Exception as e:
            print(f"\n下载失败: {e}")


if __name__ == '__main__':
    file = r'./ceshishuju.json'
    with open(file, 'r', encoding='utf-8') as file:
        data = json.load(file)

    print(data.keys())
    kzj_data = data['招标清单']
    tb_datas = data['投标文件']
    result = functions.get_qingdan_result(kzj_data=kzj_data, tb_datas=tb_datas)

    # print(len(result))
    with open('result.json', 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=4)
    # with open('result.json', 'r', encoding='utf-8') as f:
    #     result = json.load(f)
    # for name, value in result.items():
    #     print(name, len(value))
    #     for item in value:
    #         print(item)
